#!/usr/bin/env python3
"""Genera la propuesta de sponsorship en .docx para SM Homes."""
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
IMG = ROOT / "images"
OUT = ROOT / "Propuesta-Sponsorship-SM-Homes-elcanaveral.info.docx"

BRAND_BLUE = RGBColor(0x29, 0x52, 0xA3)
BRAND_DARK = RGBColor(0x0F, 0x1A, 0x35)
ACCENT_ORANGE = RGBColor(0xF9, 0x73, 0x16)
INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x6B, 0x72, 0x80)


def set_run(run, *, font="Calibri", size=11, bold=False, color=INK, italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_para(doc, text, *, size=11, bold=False, color=INK, align=None, space_after=4, font="Calibri"):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_run(r, font=font, size=size, bold=bold, color=color)
    return p


def add_h(doc, text, *, level=1):
    sizes = {1: 22, 2: 16, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, font="Calibri", size=sizes.get(level, 13), bold=True, color=BRAND_BLUE)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D5DBE5")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_bullet(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_run(r, size=11, bold=True, color=INK)
        r2 = p.add_run(" — " + text if text else "")
        set_run(r2, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run(r, size=11, color=INK)


def add_image_centered(doc, path, width_cm=15.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        set_run(cr, size=9, color=MUTED, italic=True)


def page_break(doc):
    doc.add_page_break()


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def main():
    doc = Document()

    # Margenes
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # =========================================================
    # PORTADA
    # =========================================================
    add_image_centered(doc, IMG / "cover-canaveral-twilight.jpg", width_cm=16.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("PROPUESTA DE SPONSORSHIP")
    set_run(r, font="Calibri", size=11, bold=True, color=ACCENT_ORANGE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Patrocinio exclusivo del directorio")
    set_run(r, font="Calibri", size=24, bold=True, color=BRAND_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("elcanaveral.info")
    set_run(r, font="Calibri", size=24, bold=True, color=BRAND_BLUE)

    add_divider(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Preparada para")
    set_run(r, size=10, color=MUTED, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("SM HOMES — Gestión Inmobiliaria")
    set_run(r, size=16, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Oficina El Cañaveral · Calle Enrique Urquijo 90, 28052 Madrid")
    set_run(r, size=10, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"Madrid, abril de 2026")
    set_run(r, size=10, color=MUTED)

    page_break(doc)

    # =========================================================
    # 1. RESUMEN EJECUTIVO
    # =========================================================
    add_h(doc, "Resumen ejecutivo", level=1)
    add_para(doc,
        "elcanaveral.info es el directorio hiperlocal de referencia de El Cañaveral, "
        "Vicálvaro, Coslada y San Fernando de Henares. En menos de un mes desde su lanzamiento, "
        "el directorio cuenta ya con 111 negocios, 8 inmobiliarias en su categoría dedicada, "
        "164 páginas SEO indexables y un hub multi-zona pensado para ser el primer sitio al que "
        "un vecino o un nuevo residente acude cuando busca cualquier servicio del barrio.",
        space_after=8)

    add_para(doc,
        "Más allá del directorio de comercios, nuestra ambición es ser el hub completo de El "
        "Cañaveral: negocios, inmobiliarias, perfiles sociales del barrio, servicios públicos, "
        "transporte, eventos, vida vecinal. Todo en un solo lugar.",
        space_after=10)

    add_h(doc, "Lo que proponemos", level=2)
    add_bullet(doc, "convertirse en el sponsor exclusivo del directorio durante los próximos 12 meses.",
               bold_prefix="SM Homes")
    add_bullet(doc, "tarifa de patrocinio integral con presencia destacada en home, sección Inmobiliarias, "
               "fichas y banners site-wide.", bold_prefix="800€/mes")
    add_bullet(doc, "los dos primeros meses son gratuitos mientras el directorio se asienta y "
               "consolida tráfico orgánico. Activación inmediata, sin coste hasta el mes 3.",
               bold_prefix="Oferta lanzamiento")
    add_bullet(doc, "ninguna otra inmobiliaria podrá ocupar la posición #1 ni los banners site-wide "
               "durante la vigencia del acuerdo.", bold_prefix="Exclusividad de categoría")

    add_h(doc, "Por qué SM Homes", level=2)
    add_para(doc,
        "Esta propuesta no se está extendiendo en abierto. Se está abriendo conversación con varias "
        "agencias inmobiliarias del distrito, pero hay tres motivos por los que SM Homes encaja mejor "
        "que el resto:",
        space_after=6)
    add_bullet(doc, "su oficina está en el corazón del PAU (Calle Enrique Urquijo 90), "
               "literalmente al lado del operador del directorio.", bold_prefix="Cercanía física")
    add_bullet(doc, "5★ en Google con 92 reseñas y 3 años de trayectoria visible en su escaparate.",
               bold_prefix="Reputación")
    add_bullet(doc, "el rótulo dorado, la madera natural y el lenguaje 'guiando sueños y creando "
               "oportunidades' encajan con la estética premium que estamos construyendo en el directorio.",
               bold_prefix="Estética alineada")

    page_break(doc)

    # =========================================================
    # 2. EL PROYECTO
    # =========================================================
    add_h(doc, "El proyecto: elcanaveral.info", level=1)
    add_para(doc,
        "Un directorio hiperlocal cuidado al detalle: cada ficha tiene fotos reales, datos "
        "verificados, ubicación, horarios y contacto. Diseño premium pensado para que la marca de un "
        "sponsor se vea como en una revista, no como un banner barato.",
        space_after=8)

    add_image_centered(doc, IMG / "screenshot-home-hero.jpeg", width_cm=15.5,
                       caption="Home de elcanaveral.info — vista de bienvenida")

    add_h(doc, "Cifras actuales", level=2)
    cifras = [
        ("111", "negocios verificados"),
        ("16", "categorías"),
        ("4", "zonas (El Cañaveral, Vicálvaro, Coslada, S. Fdo.)"),
        ("164", "páginas SEO indexables"),
        ("550+", "fotos reales de fachadas e interiores"),
        ("8", "inmobiliarias activas en la categoría"),
    ]
    table = doc.add_table(rows=2, cols=6)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (num, label) in enumerate(cifras):
        c1 = table.cell(0, i)
        c2 = table.cell(1, i)
        for c in (c1, c2):
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p1.add_run(num)
        set_run(r, size=22, bold=True, color=BRAND_BLUE)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        set_run(r2, size=9, color=MUTED)

    add_para(doc, "", space_after=6)

    add_h(doc, "Visión: el hub de El Cañaveral", level=2)
    add_para(doc,
        "Estamos construyendo más que un directorio. elcanaveral.info ya tiene una segunda capa "
        "—el hub Comunidad— y una hoja de ruta editorial pensada para convertirse en el primer "
        "destino digital del barrio.",
        space_after=4)
    add_bullet(doc, "directorio de perfiles sociales del barrio (Facebook, Instagram, TikTok, X) — YA LIVE.",
               bold_prefix="Comunidad")
    add_bullet(doc, "calendario de eventos vecinales y promociones locales.")
    add_bullet(doc, "guías prácticas: dónde matricular en el cole, transporte, servicios municipales.")
    add_bullet(doc, "blog editorial sobre vida en el PAU con cobertura SEO local.")
    add_bullet(doc, "newsletter mensual a vecinos suscritos al directorio.")

    page_break(doc)

    # =========================================================
    # 3. LA SECCION INMOBILIARIAS
    # =========================================================
    add_h(doc, "La sección Inmobiliarias", level=1)
    add_para(doc,
        "Hemos creado una categoría dedicada para inmobiliarias precisamente porque es el sector "
        "con mayor potencial publicitario en un PAU joven en plena consolidación. Hoy hay 8 agencias "
        "listadas distribuidas por las 4 zonas del directorio:",
        space_after=8)
    add_bullet(doc, "El Cañaveral (3): SM Homes · Mediadores Inmobiliarios · Redpiso")
    add_bullet(doc, "Vicálvaro (4): Templo Consulting · INMOACIERTA · Globalpiso · Alianza Madrid")
    add_bullet(doc, "Coslada (1): Inmobiliaria El Cañaveral")
    add_para(doc, "", space_after=6)

    add_image_centered(doc, IMG / "screenshot-categoria-inmobiliarias.jpeg", width_cm=14,
                       caption="Sección /inmobiliarias/ tal como se ve hoy")

    add_h(doc, "Qué cambia con SM Homes como sponsor", level=2)
    add_bullet(doc, "ficha pasa al puesto #1 en /inmobiliarias y /zona/el-canaveral con badge 'Sponsor oficial'.")
    add_bullet(doc, "marca SM Homes en el banner promocional superior site-wide (visible en las 164 páginas).")
    add_bullet(doc, "banner inline en cada ficha de inmobiliaria competidora ('¿buscas otra opción? prueba SM Homes').")
    add_bullet(doc, "mención fija en home: 'Inmobiliaria oficial del directorio'.")
    add_bullet(doc, "exclusividad: ninguna otra agencia podrá ocupar estas posiciones durante el contrato.")

    page_break(doc)

    # =========================================================
    # 4. LA FICHA PREMIUM
    # =========================================================
    add_h(doc, "Cómo se ve hoy la ficha de SM Homes", level=1)
    add_para(doc,
        "Antes incluso de la conversación, ya hemos preparado la ficha como muestra del nivel "
        "de cuidado del directorio. Datos verificados vía Google Places API, 5 fotos reales (fachada, "
        "interior, equipo) y CTAs directos a llamada y web.",
        space_after=8)

    add_image_centered(doc, IMG / "screenshot-ficha-sm-homes.jpeg", width_cm=15.5,
                       caption="Ficha actual: /inmobiliarias/sm-homes-el-canaveral/")

    add_para(doc,
        "Como sponsor, esta ficha incorporaría además: badge dorado 'Patrocinador oficial', "
        "tarjeta destacada en el sidebar de TODAS las fichas de la categoría con CTA directo a SM Homes, "
        "y posición fijada en el primer puesto independientemente del orden por defecto.",
        space_after=8)

    page_break(doc)

    # =========================================================
    # 5. EL HUB COMUNIDAD
    # =========================================================
    add_h(doc, "El hub Comunidad: más allá del directorio", level=1)
    add_para(doc,
        "Hemos lanzado /comunidad/, una sección que reúne a toda la conversación digital del barrio "
        "en una sola página. No existe nada parecido en El Cañaveral. Lo que conseguimos con esto: "
        "que el vecino que busca 'qué pasa en el barrio' aterrice en elcanaveral.info, no en una "
        "búsqueda dispersa por Facebook o Instagram.",
        space_after=8)

    add_h(doc, "Cifras del hub Comunidad", level=2)
    com = [
        ("75+", "cuentas verificadas"),
        ("8", "plataformas (FB, IG, X, TikTok, YouTube, Telegram, WhatsApp, web)"),
        ("16", "páginas Facebook (incluida Junta Municipal Vicálvaro oficial)"),
        ("25+", "perfiles Instagram (vecinos, comercios, asociaciones)"),
        ("6.2K", "miembros del grupo 'De El Cañaveral al cielo' indexado"),
        ("3", "niveles: Imprescindibles · Destacadas · Más cuentas"),
    ]
    t = doc.add_table(rows=2, cols=6)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (n, l) in enumerate(com):
        c1 = t.cell(0, i); c2 = t.cell(1, i)
        for c in (c1, c2):
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p1.add_run(n); set_run(r, size=18, bold=True, color=BRAND_BLUE)
        p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(l); set_run(r2, size=8, color=MUTED)
    add_para(doc, "", space_after=6)

    add_h(doc, "Por qué importa para SM Homes", level=2)
    add_bullet(doc, "el sponsor del directorio es también el sponsor del hub social del barrio. "
               "Reconocimiento de marca asociado a la 'voz oficial' de El Cañaveral.",
               bold_prefix="Asociación de marca")
    add_bullet(doc, "los vecinos que llegan vía Comunidad son residentes activos en el barrio — "
               "exactamente el público objetivo de una inmobiliaria que vende y alquila aquí.",
               bold_prefix="Tráfico cualificado")
    add_bullet(doc, "el banner site-wide y la mención en home aparecen también en Comunidad. "
               "El hub aporta tráfico extra al inventario publicitario del sponsor.",
               bold_prefix="Inventario adicional")

    page_break(doc)

    # =========================================================
    # 6. PROYECCION DE TRAFICO
    # =========================================================
    add_h(doc, "Proyección de tráfico y crecimiento", level=1)
    add_para(doc,
        "Estimaciones basadas en el ritmo de indexación actual, la velocidad de incorporación de "
        "negocios y el plan editorial de los próximos 9 meses. SEO local en un PAU joven con poca "
        "competencia es una de las jugadas con mejor retorno temporal del marketing digital.",
        space_after=8)

    proj = doc.add_table(rows=6, cols=4)
    proj.style = "Light Grid Accent 1"
    proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headers = ["Periodo", "Visitas/mes (est.)", "Páginas vistas", "Hito"]
    rows = [
        ("Q2 2026 (actual)", "1.500 – 3.000", "5.000 – 9.000", "Lanzamiento, indexación inicial"),
        ("Q3 2026", "6.000 – 10.000", "20.000 – 35.000", "Pilares SEO + 200 negocios"),
        ("Q4 2026", "15.000 – 22.000", "50.000 – 75.000", "Blog editorial + newsletter activa"),
        ("Q1 2027", "25.000 – 35.000", "85.000 – 120.000", "Marca instalada en el PAU"),
        ("Cierre 2027 (est.)", "45.000 – 60.000", "150.000 – 200.000", "Referente del distrito Vicálvaro"),
    ]
    for j, h in enumerate(headers):
        c = proj.cell(0, j)
        shade_cell(c, "1F2937")
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); set_run(r, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = proj.cell(i, j)
            p = c.paragraphs[0]
            if j > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val); set_run(r, size=10, color=INK, bold=(j == 0))
    add_para(doc, "", space_after=6)

    add_para(doc,
        "La oferta de meses 1-2 gratuitos de SM Homes coincide exactamente con el tramo de mayor "
        "ROI del sponsor: entrar antes de la curva de crecimiento, no después.",
        size=10, color=MUTED, space_after=8)

    page_break(doc)

    # =========================================================
    # 7. PALABRAS CLAVE QUE ATACAMOS
    # =========================================================
    add_h(doc, "Palabras clave que estamos atacando", level=1)
    add_para(doc,
        "El SEO del directorio se ha estructurado como una matriz de servicio × zona: cada categoría "
        "se cruza con cada zona y cada subzona, generando cientos de páginas indexables que capturan "
        "intención de búsqueda hiperlocal. Estos son los grupos de keywords con mayor prioridad:",
        space_after=8)

    add_h(doc, "Inmobiliario (prioridad sponsor)", level=2)
    add_bullet(doc, "inmobiliaria El Cañaveral · agencia inmobiliaria El Cañaveral")
    add_bullet(doc, "pisos en venta El Cañaveral · pisos alquiler PAU Cañaveral")
    add_bullet(doc, "obra nueva Cañaveral · viviendas Vicálvaro · pisos Coslada")
    add_bullet(doc, "inmobiliaria PAU Cañaveral · vender piso Cañaveral")
    add_bullet(doc, "comprar casa El Cañaveral Madrid · alquiler PAU Cañaveral")

    add_h(doc, "Servicios diarios", level=2)
    add_bullet(doc, "supermercados El Cañaveral · fruterías Cañaveral · panaderías PAU")
    add_bullet(doc, "restaurantes El Cañaveral · cafeterías Cañaveral · pizzerías PAU")
    add_bullet(doc, "farmacia El Cañaveral · dentista Cañaveral · clínica PAU")
    add_bullet(doc, "peluquería El Cañaveral · barbería Cañaveral · estética Vicálvaro")

    add_h(doc, "Familias y mascotas", level=2)
    add_bullet(doc, "guardería El Cañaveral · colegio CEIPSO Rudyard Kipling")
    add_bullet(doc, "academia inglés Cañaveral · academia repaso PAU")
    add_bullet(doc, "veterinario El Cañaveral · peluquería canina Cañaveral")

    add_h(doc, "Hogar, deporte y ocio", level=2)
    add_bullet(doc, "ferretería El Cañaveral · fontanero Cañaveral · electricista PAU")
    add_bullet(doc, "gimnasio El Cañaveral · pádel Cañaveral · pilates PAU")
    add_bullet(doc, "talleres mecánicos Cañaveral · ITV Vicálvaro")

    add_h(doc, "Búsquedas de comunidad y vida en el barrio", level=2)
    add_bullet(doc, "qué hacer en El Cañaveral · eventos PAU Cañaveral")
    add_bullet(doc, "asociación vecinos El Cañaveral · noticias El Cañaveral")
    add_bullet(doc, "Junta Municipal Vicálvaro · transporte PAU Cañaveral")
    add_bullet(doc, "vivir en El Cañaveral · barrio El Cañaveral Madrid")

    page_break(doc)

    # =========================================================
    # 8. HOJA DE RUTA
    # =========================================================
    add_h(doc, "Hoja de ruta: lo que viene", level=1)
    add_para(doc,
        "El sponsorship de SM Homes financia la consolidación de elcanaveral.info como hub completo "
        "del barrio. Estas son las áreas que activamos en los próximos 12 meses, todas ellas con "
        "presencia visible del sponsor:",
        space_after=8)

    add_h(doc, "Q3 2026 — Consolidación", level=2)
    add_bullet(doc, "ampliación a 200+ negocios (cobertura 95% comercio físico del PAU).",
               bold_prefix="Directorio")
    add_bullet(doc, "calendario público de eventos del barrio (mercados, fiestas, asociaciones).",
               bold_prefix="Eventos")
    add_bullet(doc, "guías editoriales: 'cómo matricular en el cole', 'transporte al centro', "
               "'servicios municipales'.", bold_prefix="Guías prácticas")

    add_h(doc, "Q4 2026 — Editorial y Audiencia", level=2)
    add_bullet(doc, "blog con 1-2 piezas semanales sobre vida en el PAU. Cobertura SEO long-tail.",
               bold_prefix="Blog")
    add_bullet(doc, "boletín mensual con novedades del barrio y nuevos negocios listados.",
               bold_prefix="Newsletter")
    add_bullet(doc, "mapa visual de El Cañaveral con todos los servicios geolocalizados.",
               bold_prefix="Mapa interactivo")

    add_h(doc, "Q1 2027 — Servicios al Vecino", level=2)
    add_bullet(doc, "información actualizada de líneas de bus, EMT, próximas estaciones de Metro.",
               bold_prefix="Transporte")
    add_bullet(doc, "directorio de centros de salud, especialistas concertados, farmacias de guardia.",
               bold_prefix="Sanidad")
    add_bullet(doc, "información oficial de la Junta Municipal de Vicálvaro y Ayuntamiento.",
               bold_prefix="Servicios públicos")

    add_h(doc, "Q2 2027 — Marca de Barrio", level=2)
    add_bullet(doc, "merchandising/identidad asociada a 'Hecho en El Cañaveral'.",
               bold_prefix="Branding del barrio")
    add_bullet(doc, "patrocinios cruzados con asociaciones vecinales y eventos locales.",
               bold_prefix="Comunidad")
    add_bullet(doc, "extensión a otros PAUs de Madrid bajo la misma plataforma LocalSEOAds.",
               bold_prefix="Expansión")

    page_break(doc)

    # =========================================================
    # 9. INVERSION Y TERMINOS
    # =========================================================
    add_h(doc, "Inversión y términos", level=1)

    # Tabla de inversion
    inv = doc.add_table(rows=5, cols=2)
    inv.style = "Light List Accent 1"
    inv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rows = [
        ("Tarifa estándar", "800 €/mes (9.600 € anuales)"),
        ("Oferta de lanzamiento", "Meses 1 y 2 GRATUITOS"),
        ("Facturación efectiva mes 3-12", "800 € × 10 = 8.000 €"),
        ("Total año 1", "8.000 € (ahorro de 1.600 € sobre tarifa)"),
        ("Permanencia", "Contrato anual con renovación automática"),
    ]
    for i, (k, v) in enumerate(rows):
        c1 = inv.cell(i, 0); c2 = inv.cell(i, 1)
        c1.width = Cm(7); c2.width = Cm(8)
        r1 = c1.paragraphs[0].add_run(k)
        set_run(r1, size=11, bold=True, color=INK)
        r2 = c2.paragraphs[0].add_run(v)
        set_run(r2, size=11, color=INK)
    add_para(doc, "", space_after=6)

    add_h(doc, "Qué incluye", level=2)
    add_bullet(doc, "TopBanner site-wide (presente en las 164 páginas indexadas).", bold_prefix="Visibilidad")
    add_bullet(doc, "ficha #1 en /inmobiliarias y /zona/el-canaveral con badge 'Patrocinador oficial'.",
               bold_prefix="Posicionamiento")
    add_bullet(doc, "tarjeta SM Homes en el sidebar de TODAS las fichas de inmobiliarias competidoras.",
               bold_prefix="Cross-sell")
    add_bullet(doc, "mención en home: 'Inmobiliaria oficial del directorio'.", bold_prefix="Reconocimiento")
    add_bullet(doc, "ninguna otra agencia ocupa las mismas posiciones durante el contrato.",
               bold_prefix="Exclusividad")
    add_bullet(doc, "informe mensual de tráfico: páginas vistas, clics a la ficha SM Homes, "
               "llamadas desde la ficha.", bold_prefix="Reporte")
    add_bullet(doc, "subida de listado de propiedades destacadas si en algún momento se quiere "
               "ampliar el ámbito (sin coste adicional el primer año).",
               bold_prefix="Bonus")

    page_break(doc)

    # =========================================================
    # 6. PROXIMOS PASOS
    # =========================================================
    add_h(doc, "Próximos pasos", level=1)

    pasos = [
        ("1", "Conversación inicial",
         "Esta propuesta se entrega en mano. Resolvemos dudas y ajustamos detalles."),
        ("2", "Firma del acuerdo",
         "Documento simple, contrato anual con cláusula de salida tras los 12 meses."),
        ("3", "Activación inmediata",
         "Aplicamos featured + badge + banners en menos de 24 horas tras la firma."),
        ("4", "Reporte mensual",
         "Cada mes recibe SM Homes un informe de tráfico, leads y rendimiento."),
        ("5", "Renovación / ampliación",
         "Al cierre del año 1 evaluamos resultados y ofrecemos renovación con condiciones preferentes."),
    ]
    for num, titulo, desc in pasos:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"  {num}.  ")
        set_run(r1, size=12, bold=True, color=ACCENT_ORANGE)
        r2 = p.add_run(titulo)
        set_run(r2, size=12, bold=True, color=INK)
        d = doc.add_paragraph()
        d.paragraph_format.left_indent = Cm(0.9)
        d.paragraph_format.space_after = Pt(8)
        rd = d.add_run(desc)
        set_run(rd, size=10, color=MUTED)

    add_divider(doc)

    add_h(doc, "Contacto", level=2)
    add_para(doc, "John Guerrero", size=12, bold=True, color=INK, space_after=2)
    add_para(doc, "Propietario de LocalSEOAds.com — titular oficial de elcanaveral.info",
             size=10, color=MUTED, space_after=8)

    add_para(doc, "Para esta propuesta y temas del directorio:", size=10, bold=True, color=INK, space_after=2)
    add_para(doc, "local@elcanaveral.info", size=10, color=BRAND_BLUE, space_after=2)
    add_para(doc, "www.elcanaveral.info", size=10, color=MUTED, space_after=8)

    add_para(doc, "Para servicios de agencia digital (publicidad, web, SEO de tu empresa):",
             size=10, bold=True, color=INK, space_after=2)
    add_para(doc, "info@localseoads.com", size=10, color=BRAND_BLUE, space_after=2)
    add_para(doc, "www.localseoads.com", size=10, color=MUTED, space_after=12)

    add_divider(doc)

    add_h(doc, "Firma", level=2)
    add_para(doc,
        "Esta propuesta de sponsorship la firma LocalSEOAds.com, en calidad de titular y operador "
        "del directorio elcanaveral.info, representada por John Guerrero como propietario oficial "
        "de ambas entidades.",
        size=10, color=INK, space_after=10)

    add_para(doc, "John Guerrero", size=12, bold=True, color=INK, space_after=0)
    add_para(doc, "Propietario · LocalSEOAds.com",
             size=10, color=MUTED, space_after=0)
    add_para(doc, "Titular del directorio elcanaveral.info",
             size=10, color=MUTED, space_after=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("Gracias por considerar la propuesta.")
    set_run(r, size=11, italic=True, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Construyamos juntos el barrio.")
    set_run(r, size=12, bold=True, color=BRAND_BLUE)

    doc.save(OUT)
    print(f"Guardado: {OUT}")
    print(f"Tamaño: {OUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
